"""Yerel bilgi tabanındaki desteklenen belgeleri güvenle okur ve işler.

Bu modül Markdown (.md), Metin (.txt), PDF (.pdf), Word (.docx), Excel (.xlsx) ve CSV (.csv)
dosyalarını okur, metin içeriğini çıkarır, normalleştirir ve OCR fallback desteği sunar.
"""

from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from app.config import SUPPORTED_SUFFIXES
from app.domain import SourceDocument


def normalize_text(text: str) -> str:
    """Metindeki satır sonu karakterlerini (CRLF, CR) Unix formatına (\n) dönüştürür ve kenar boşlukları temizler.

    Args:
        text (str): İşlenecek ham metin.

    Returns:
        str: Formatı düzeltilmiş ve temizlenmiş metin.
    """
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _read_file_content(path: Path) -> str:
    """Farklı uzantılara sahip dosyaların türüne uygun kütüphaneler ile metin içeriğini çıkarır.

    Args:
        path (Path): Okunacak dosyanın tam veya bağıl yolu.

    Returns:
        str: Çıkarılan metin içeriği.

    Raises:
        ValueError: İlgili dosya türünü okumak için gerekli bağımlılık yoksa veya dosya desteklenmiyorsa.
    """
    suffix = path.suffix.lower()

    # Düz metin ve Markdown belgelerinin doğrudan okunması
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8")

    # PDF belgelerinin işlenmesi ve gerektiğinde Tesseract OCR çalıştırılması
    elif suffix == ".pdf":
        try:
            import pypdf  # type: ignore
            reader = pypdf.PdfReader(path)
            pages_text: list[str] = []
            ocr_needed_pages: list[int] = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(text)
                else:
                    ocr_needed_pages.append(i)
                    pages_text.append("")  # OCR için yer tutucu

            # Taranmış veya metin içermeyen sayfalar için OCR fallback
            if ocr_needed_pages:
                try:
                    from pdf2image import convert_from_path  # type: ignore
                    import pytesseract  # type: ignore

                    if shutil.which("tesseract") is None:
                        raise ValueError("Bu PDF taranmış görünüyor. OCR için Tesseract kurulmalı ve PATH'e eklenmeli.")
                    try:
                        pytesseract.get_tesseract_version()
                    except Exception as exc:
                        raise ValueError("Tesseract çalıştırılamadı. Tesseract ve tur.traineddata kurulumunu kontrol edin.") from exc
                    images = convert_from_path(str(path))
                    for page_idx in ocr_needed_pages:
                        if page_idx < len(images):
                            ocr_text = pytesseract.image_to_string(images[page_idx], lang="tur+eng")
                            pages_text[page_idx] = ocr_text
                except ImportError as exc:
                    raise ValueError("Taranmış PDF için OCR bağımlılıkları eksik: pip install pytesseract pdf2image") from exc

            return "\n".join(pages_text)
        except ImportError:
            raise ValueError(f"PDF dosyasını okumak için 'pypdf' kütüphanesi kurulu olmalıdır: {path}")

    # Word (.docx) belgelerinin okunması
    elif suffix == ".docx":
        try:
            import docx  # type: ignore
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text)
        except ImportError:
            raise ValueError(f"DOCX dosyasını okumak için 'python-docx' kütüphanesi kurulu olmalıdır: {path}")

    # Tablo tabanlı CSV belgelerinin sütun isimleriyle birlikte okunması
    elif suffix == ".csv":
        import csv
        lines: list[str] = []
        with open(path, encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            header: list[str] | None = None
            for row in reader:
                cleaned_row = [cell.strip() for cell in row]
                if not any(cleaned_row):
                    continue
                if header is None:
                    header = cleaned_row
                    lines.append(" | ".join(header))
                else:
                    row_pairs = [
                        f"{h}: {val}" if h else val
                        for h, val in zip(header, cleaned_row)
                        if val
                    ]
                    lines.append(" | ".join(row_pairs) if row_pairs else " | ".join(cleaned_row))
        return "\n".join(lines)

    # Excel (.xlsx) belgelerinin sayfa sayfa okunması
    elif suffix == ".xlsx":
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            try:
                sheets_text: list[str] = []
                for sheet in wb.worksheets:
                    sheets_text.append(f"Sayfa: {sheet.title}")
                    for row in sheet.iter_rows(values_only=True):
                        row_vals = [str(val).strip() for val in row if val is not None and str(val).strip()]
                        if row_vals:
                            sheets_text.append(" | ".join(row_vals))
                return "\n".join(sheets_text)
            finally:
                wb.close()
        except ImportError:
            raise ValueError(f"XLSX dosyasını okumak için 'openpyxl' kütüphanesi kurulu olmalıdır: {path}")

    raise ValueError(f"Desteklenmeyen dosya biçimi: {path}")


def discover_documents(knowledge_base_dir: Path) -> list[SourceDocument]:
    """Bilgi tabanı dizinindeki desteklenen tüm dosyaları tarar ve SourceDocument nesneleri olarak yükler.

    Args:
        knowledge_base_dir (Path): Taranacak ana bilgi tabanı dizini.

    Returns:
        list[SourceDocument]: Keşfedilen geçerli ve dolu kaynak belgelerin listesi.

    Raises:
        ValueError: Dosya okuma veya UTF-8 kodlama hatası yaşanırsa.
    """
    if not knowledge_base_dir.exists():
        return []

    documents: list[SourceDocument] = []
    # Dosyaları alfabetik sırayla tarama
    for path in sorted(knowledge_base_dir.rglob("*"), key=lambda item: item.as_posix().lower()):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue

        try:
            raw_content = _read_file_content(path)
            content = normalize_text(raw_content)
        except UnicodeDecodeError as exc:
            raise ValueError(f"Belge UTF-8 olarak okunamadı: {path}") from exc
        except Exception as exc:
            raise ValueError(f"Belge okuma hatası: {path}") from exc

        if not content:
            continue

        source_path = path.relative_to(knowledge_base_dir.parent).as_posix()
        content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
        documents.append(SourceDocument(source_path, content, content_hash))

    return documents

